import os
import spacy
import fitz  # PyMuPDF
import base64
import requests  # For downloading files
import time  # For timing
from docx import Document
from flask import Flask, request, jsonify

# Initialize Flask app
app = Flask(__name__)

# Load spaCy's English model
nlp = spacy.load("en_core_web_sm")

# API keys array
API_KEYS = ["111111111111111", "222222222222222", "333333333333333"]


@app.before_request
def validate_api_key():
    """Validate API key in the headers."""
    api_key = request.headers.get("API-Key")
    if api_key not in API_KEYS:
        return jsonify({"error": "Invalid or missing API key"}), 403


def clean_and_reduce_text(text):
    """Reduce text by removing unnecessary details and filtering meaningful sentences."""
    doc = nlp(text)
    reduced_sentences = [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 20]
    reduced_text = " ".join(reduced_sentences)
    return reduced_text, len(doc), len(reduced_sentences)


def extract_text_from_pdf(pdf_path, reduce_tokens=False, include_images=True):
    """Extract text and optionally images from a PDF file."""
    start_time = time.time()
    doc = fitz.Document(pdf_path)
    text = ""
    images = []

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        page_text = page.get_text().strip()
        text += page_text + " "

        if include_images:
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                base64_image = base64.b64encode(image_bytes).decode("utf-8")
                images.append({"image_index": img_index + 1, "base64": base64_image})

    image_count = len(images) if include_images else 0

    if reduce_tokens:
        reduced_text, token_count, sentence_count = clean_and_reduce_text(text)
    else:
        reduced_text, token_count, sentence_count = text, len(nlp(text)), len(list(nlp(text).sents))

    process_time = time.time() - start_time
    file_size = os.path.getsize(pdf_path)

    result = {
        "reduced_text": reduced_text,
        "full_text": text,
        "token_count": token_count,
        "sentence_count": sentence_count,
        "process_time_seconds": process_time,
        "file_size_bytes": file_size,
    }

    if include_images:
        result["image_count"] = image_count
        result["images"] = images

    return result


def extract_text_from_docx(docx_path, reduce_tokens=False, include_images=True):
    """Extract text and optionally images from a DOCX file."""
    start_time = time.time()
    doc = Document(docx_path)
    text = ""
    images = []

    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:
            text += para_text + " "

    if include_images:
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                base64_image = base64.b64encode(image_data).decode("utf-8")
                images.append({"image_index": len(images) + 1, "base64": base64_image})

    image_count = len(images) if include_images else 0

    if reduce_tokens:
        reduced_text, token_count, sentence_count = clean_and_reduce_text(text)
    else:
        reduced_text, token_count, sentence_count = text, len(nlp(text)), len(list(nlp(text).sents))

    process_time = time.time() - start_time
    file_size = os.path.getsize(docx_path)

    result = {
        "reduced_text": reduced_text,
        "full_text": text,
        "token_count": token_count,
        "sentence_count": sentence_count,
        "process_time_seconds": process_time,
        "file_size_bytes": file_size,
    }

    if include_images:
        result["image_count"] = image_count
        result["images"] = images

    return result


def process_pdf_from_base64(encoded_data, reduce_tokens=False, include_images=True):
    """Process PDF file from Base64 encoded data."""
    file_path = "/tmp/temp.pdf"
    with open(file_path, "wb") as f:
        f.write(base64.b64decode(encoded_data))

    result = extract_text_from_pdf(file_path, reduce_tokens, include_images)
    os.remove(file_path)
    return result


def process_docx_from_base64(encoded_data, reduce_tokens=False, include_images=True):
    """Process DOCX file from Base64 encoded data."""
    file_path = "/tmp/temp.docx"
    with open(file_path, "wb") as f:
        f.write(base64.b64decode(encoded_data))

    result = extract_text_from_docx(file_path, reduce_tokens, include_images)
    os.remove(file_path)
    return result


@app.route("/process_pdf_url", methods=["POST"])
def process_pdf_url():
    """Process a PDF file from a URL."""
    data = request.get_json()
    file_url = data.get("file_url")
    reduce_tokens = data.get("reduce_tokens", False)
    include_images = data.get("include_images", True)

    if not file_url:
        return jsonify({"error": "file_url is required"}), 400

    response = requests.get(file_url, stream=True)
    if response.status_code != 200:
        return jsonify({"error": "Failed to download file."}), 400

    file_path = "/tmp/temp.pdf"
    with open(file_path, "wb") as f:
        f.write(response.content)

    result = extract_text_from_pdf(file_path, reduce_tokens, include_images)
    os.remove(file_path)
    return jsonify(result)


@app.route("/process_docx_url", methods=["POST"])
def process_docx_url():
    """Process a DOCX file from a URL."""
    data = request.get_json()
    file_url = data.get("file_url")
    reduce_tokens = data.get("reduce_tokens", False)
    include_images = data.get("include_images", True)

    if not file_url:
        return jsonify({"error": "file_url is required"}), 400

    response = requests.get(file_url, stream=True)
    if response.status_code != 200:
        return jsonify({"error": "Failed to download file."}), 400

    file_path = "/tmp/temp.docx"
    with open(file_path, "wb") as f:
        f.write(response.content)

    result = extract_text_from_docx(file_path, reduce_tokens, include_images)
    os.remove(file_path)
    return jsonify(result)


@app.route("/process_pdf_base64", methods=["POST"])
def process_pdf_base64():
    """Process a PDF file from Base64 encoded data."""
    data = request.get_json()
    encoded_data = data.get("file_base64")
    reduce_tokens = data.get("reduce_tokens", False)
    include_images = data.get("include_images", True)

    if not encoded_data:
        return jsonify({"error": "file_base64 is required"}), 400

    result = process_pdf_from_base64(encoded_data, reduce_tokens, include_images)
    return jsonify(result)


@app.route("/process_docx_base64", methods=["POST"])
def process_docx_base64():
    """Process a DOCX file from Base64 encoded data."""
    data = request.get_json()
    encoded_data = data.get("file_base64")
    reduce_tokens = data.get("reduce_tokens", False)
    include_images = data.get("include_images", True)

    if not encoded_data:
        return jsonify({"error": "file_base64 is required"}), 400

    result = process_docx_from_base64(encoded_data, reduce_tokens, include_images)
    return jsonify(result)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)
